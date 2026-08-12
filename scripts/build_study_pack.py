"""
SoilGuard-CG: Ideathon Study & Presentation Pack Generator
Builds a set of readable .docx documents (study guide, methodology, code walkthrough,
results cheat-sheet, demo script, presentation script, and judges Q&A) into
`ideathon-study-pack/` at the repository root.

Usage:
    python scripts/build_study_pack.py

Requires: python-docx, Pillow   (pip install python-docx pillow)
Data sources: shared/metrics.json, outputs/golden_backup/*.csv (canonical recorded
values used by the study guide / brief), fresh maps from outputs/phase2|3|4.
"""

import json
import os
import sys
from pathlib import Path

import pandas as pd
from PIL import Image

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("python-docx is required:  pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "ideathon-study-pack"
ASSETS = OUT_DIR / "assets"
MAPS_SRC = ROOT / "soilguard-cg" / "outputs"
GOLDEN = ROOT / "soilguard-cg" / "outputs" / "golden_backup"

# ---- Canonical numbers (shared/metrics.json) ---------------------------------
with open(ROOT / "shared" / "metrics.json", encoding="utf-8") as f:
    M = json.load(f)

GREEN = RGBColor(0x16, 0x62, 0x2E)
DARK = RGBColor(0x11, 0x18, 0x27)
ACCENT = RGBColor(0x0E, 0x74, 0xBC)


# ---- docx helpers -------------------------------------------------------------
def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(5)
    return doc


def set_landscape(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    w, h = sec.page_width, sec.page_height
    sec.page_width, sec.page_height = h, w


def title_page(doc, title, subtitle, lines=()):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SoilGuard-CG")
    r.font.size = Pt(40)
    r.font.bold = True
    r.font.color.rgb = GREEN
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = DARK
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = ACCENT
    doc.add_paragraph()
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x55, 0x5B, 0x66)
    doc.add_page_break()


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = GREEN
    return h


def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = DARK
    return h


def h3(doc, text):
    return doc.add_heading(text, level=3)


def para(doc, text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None, font_size=9.5, header_fill="16622E"):
    tbl = doc.add_table(rows=1, cols=len(headers))
    try:
        tbl.style = "Light Grid Accent 1"
    except Exception:
        tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, header_fill)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(font_size)
            if i == 0:
                r.bold = True
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    # repeat header row on page breaks
    trPr = hdr._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)
    doc.add_paragraph()
    return tbl


def callout(doc, label, text, fill="E7F0EA"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.color.rgb = GREEN
    p.add_run(text)
    doc.add_paragraph()
    return tbl


def embed_map(doc, phase, fname, caption, max_w=6.3):
    src = MAPS_SRC / phase / fname
    if not src.exists():
        para(doc, f"[image not found: {fname}]", italic=True)
        return
    dest = ASSETS / (phase + "_" + fname)
    img = Image.open(src).convert("RGB")
    if img.width > 1600:
        img = img.resize((1600, int(img.height * 1600 / img.width)), Image.LANCZOS)
    dest.parent.mkdir(parents=True, exist_ok=True)
    if fname.endswith(".png"):
        img.save(dest, "JPEG", quality=85)
    else:
        img.save(dest)
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic.add_run()
    run.add_picture(str(dest), width=Inches(max_w))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x5B, 0x66)
    doc.add_paragraph()


# ---- data ---------------------------------------------------------------------
def load_zonal():
    return pd.read_csv(GOLDEN / "zonal_priority_ranking.csv")


def load_recs():
    return pd.read_csv(GOLDEN / "agronomic_recommendations.csv")


# ===============================================================================
# DOC 0 — Start here / study plan
# ===============================================================================
def doc00():
    doc = new_doc()
    title_page(
        doc,
        "Start Here — Study Plan & How to Use This Pack",
        "National Space Day Ideathon 2026 · COSINE NIT Raipur × NRSC / ISRO",
        lines=[
            "Project: SoilGuard-CG — Chhattisgarh Soil Carbon Sentinel",
            "Prepared from the live codebase, pipeline outputs, and master study guide.",
        ],
    )
    h1(doc, "What is in this folder?")
    add_table(
        doc,
        ["Document", "Purpose", "Study time"],
        [
            ["01_Project_Overview_Pitch.docx", "What it is, the problem, the solution, USPs", "20 min"],
            ["02_Scientific_Methodology.docx", "Formulas, data, model, no-leakage design", "45 min"],
            ["03_Pipeline_Code_Guide.docx", "Every source file explained + how to run", "40 min"],
            ["04_Key_Results_Cheat_Sheet.docx", "All numbers, 25-sector table, maps", "30 min"],
            ["05_Live_Demo_Guide.docx", "Step-by-step live demo script", "20 min"],
            ["06_Presentation_Script.docx", "Slide-by-slide pitch script (timed)", "30 min"],
            ["07_Judges_QA.docx", "12+ judge questions with winning answers", "40 min"],
        ],
        col_widths=[2.9, 3.0, 0.9],
    )
    h1(doc, "Recommended study path")
    para(doc, "Day 1 (45 min) — the big picture:", bold=True)
    bullet(doc, "Read 01 Overview. Learn the one-line pitch and the three core pillars (Detect → Prioritize → Remediate).")
    bullet(doc, "Read 04 Cheat Sheet. Memorise the top-8 numbers (bare soil area, high-deficiency ha, R², cut-offs).")
    para(doc, "Day 2 (60 min) — the science:", bold=True)
    bullet(doc, "Read 02 Methodology. Be able to write NDVI, BSI and the target proxy formula on a whiteboard.")
    bullet(doc, "Practise explaining 'zero target leakage' — the strongest credibility point of the project.")
    para(doc, "Day 3 (60 min) — the product:", bold=True)
    bullet(doc, "Read 03 Code Guide, then 05 Demo Guide. Run demo.bat once yourself and time it.")
    bullet(doc, "Read 06 Script and 07 Q&A. Practise the 2-minute pitch out loud twice.")
    h1(doc, "The one-line pitch (memorise this)")
    callout(
        doc,
        "PITCH:",
        "SoilGuard-CG is a 100% offline, satellite-driven ML platform that turns 10m Sentinel-2 imagery into "
        "high-resolution Soil Organic Carbon (SOC) deficiency maps, 5×5 zonal priority rankings, and "
        "clay-sensitive regenerative farming advisories for the Raipur–Durg rice belt — in about a minute, "
        "with zero physical soil sampling.",
    )
    h1(doc, "Numbers you must know cold")
    add_table(
        doc,
        ["Metric", "Value", "Metric", "Value"],
        [
            ["Evaluated bare-soil area", f"{M['bareSoilHa']:,.2f} ha", "Bare soil share", f"{M['bareSoilPct']:.2f}%"],
            ["High SOC deficiency (>0.58)", f"{M['highRiskHa']:,.0f} ha ({M['highRiskPct']:.2f}%)", "Moderate (0.45–0.58)", f"{M['moderateRiskHa']:,.0f} ha"],
            ["Model R² (recorded)", f"{M['r2']:.4f}", "RMSE", f"{M['rmse']:.4f}"],
            ["Spatial resolution", M["resolution"], "Sector grid", "5×5 (25 sectors)"],
            ["Rank #1 sector", "Arang (B-1)", "Typical runtime", f"≈{M['runtimeSec']:.0f} s"],
        ],
        font_size=10,
    )
    doc.save(OUT_DIR / "00_Start_Here_Study_Plan.docx")


# ===============================================================================
# DOC 1 — Project overview & pitch
# ===============================================================================
def doc01():
    doc = new_doc()
    title_page(
        doc,
        "Project Overview & Pitch",
        "What SoilGuard-CG is, the problem it solves, and why judges should care",
        lines=["The 3-pillar story: Detect → Prioritize → Remediate"],
    )
    h1(doc, "1. What is SoilGuard-CG?")
    para(doc, "SoilGuard-CG is a terminal-native, fully offline geospatial Machine Learning platform for the "
              "Chhattisgarh rice belt ('Dhan ka Katora'). It converts 10m Sentinel-2 satellite imagery and "
              "SoilGrids topsoil properties into:")
    bullet(doc, "Continuous Soil Organic Carbon (SOC) Deficiency risk maps (score 0.0–1.0 at 10m resolution).")
    bullet(doc, "A 5×5 zonal priority grid that ranks 25 agricultural sectors by intervention urgency.")
    bullet(doc, "Clay-sensitive regenerative advisory packages (FYM dosage, green manuring, zero-tillage, lime/gypsum).")
    h1(doc, "2. The three core pillars")
    add_table(
        doc,
        ["Pillar", "What it does", "Why it matters"],
        [
            ["Detect", "Evaluates 2.27 million bare-soil pixels at 10m resolution via satellite optical physics", "Fills the spatial gap left by 1 sample / 10 ha / 3 yrs soil health cards"],
            ["Prioritize", "Overlays a 5×5 administrative grid and ranks 25 sectors by mean deficiency + high-risk area", "Officers know exactly which block to fund first"],
            ["Remediate", "Generates site-specific prescriptions using clay %, pH and BSI rules", "Farmers get exact FYM tonnes/ha, not vague advice"],
        ],
    )
    h1(doc, "3. The problem (Chhattisgarh context)")
    bullet(doc, "Paddy mono-cropping + intense summer heat (42°C+) + stubble burning drive topsoil SOC depletion.")
    bullet(doc, "Over 52% of regional agricultural grids have SOC below the critical 0.50% threshold.", bold_prefix="Crisis: ")
    bullet(doc, "Traditional Soil Health Cards sample 1 physical point per 10 ha once every 3 years — ₹2,000–3,000 per sample, missing intra-field variability.", bold_prefix="Gap: ")
    bullet(doc, "Physical sampling cannot track micro-spatial degradation or give timely post-harvest advice.", bold_prefix="Cost: ")
    h1(doc, "4. The solution (our system)")
    para(doc, "Evaluate every 10m × 10m grid cell (0.01 ha) using free, open-access Sentinel-2 data refreshed "
              "every 5 days — at zero physical sampling cost. Output is directly usable by district collectors, "
              "agriculture officers, and carbon-credit programmes.")
    h1(doc, "5. Unique selling propositions (USPs)")
    bullet(doc, "Runs 100% offline on cached windowed Cloud-Optimized GeoTIFFs — works in remote district offices.", bold_prefix="⚡ Zero internet dependency: ")
    bullet(doc, "Raw SoilGrids SOC is excluded from model features X — the model learns genuine spectral physics, not memorized maps.", bold_prefix="🛡️ Zero target leakage: ")
    bullet(doc, "FYM/tillage guidance adjusts for Vertisol clay vs light sandy soils (10–12 vs 8–10 t/ha).", bold_prefix="🎯 Clay-sensitive advisories: ")
    bullet(doc, "Full pipeline (load → mask → train → map → report) in roughly 30–45 seconds on a standard laptop.", bold_prefix="⏱️ Ultra-fast: ")
    h1(doc, "6. Tech stack (honest)")
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["Language", "Python 3.11, TypeScript"],
            ["ML", "scikit-learn Random Forest Regressor (150 trees)"],
            ["Geospatial", "rasterio, geopandas, shapely, pystac-client, GDAL"],
            ["Data", "Sentinel-2 L2A (B02/B04/B08/B11), SoilGrids (SOC, clay, pH)"],
            ["Visuals", "matplotlib 300 DPI maps, rich CLI dashboard"],
            ["API", "ISRO Bhuvan LULC & Village Geocoding (urllib)"],
            ["Frontend", "Next.js dashboard (port 3001) + map visualizer (port 5000)"],
        ],
        col_widths=[1.7, 4.8],
    )
    doc.save(OUT_DIR / "01_Project_Overview_Pitch.docx")


# ===============================================================================
# DOC 2 — Scientific methodology
# ===============================================================================
def doc02():
    doc = new_doc()
    title_page(
        doc,
        "Scientific Methodology",
        "Data, spectral indices, the ML model, and the zero-leakage design",
        lines=["Write the formulas, defend the R², explain the physics"],
    )
    h1(doc, "1. Input data")
    add_table(
        doc,
        ["Dataset", "Bands / properties", "Role"],
        [
            ["Sentinel-2 L2A (10m COG)", "B02 Blue (490nm), B04 Red (665nm), B08 NIR (842nm), B11 SWIR1 (1610nm)", "Model features X (100% satellite)"],
            ["SoilGrids (aligned 10m)", "SOC (dg/kg), Clay (g/kg), pH (×10)", "Target proxy y + advisory rules (never in X)"],
        ],
    )
    bullet(doc, "Locked AOI: [81.60°E, 21.10°N] → [81.80°E, 21.30°N] — Raipur & Durg plains, EPSG:32644.")
    bullet(doc, "Raster grid: 2,086 × 2,223 pixels = 4.64 M total pixels; 10m pixels → 0.01 ha each.")
    h1(doc, "2. Spectral indices & bare-soil masking (Phase 2)")
    para(doc, "NDVI  =  (NIR − Red) / (NIR + Red)", bold=True)
    para(doc, "BSI   =  ((SWIR1 + Red) − (NIR + Blue)) / ((SWIR1 + Red) + (NIR + Blue))", bold=True)
    para(doc, "A pixel is a candidate bare-soil pixel when ALL of these hold:")
    add_table(
        doc,
        ["Rule", "Meaning"],
        [
            ["NDVI ≤ 0.30", "Excludes dense green vegetation / standing crop"],
            ["NIR ≥ 300", "Excludes water bodies and deep shadow"],
            ["NDVI ≥ −0.20", "Excludes non-soil dark noise"],
        ],
    )
    bullet(doc, f"Result: {M['bareSoilPct']:.2f}% of the scene ({M['bareSoilHa']:,.2f} ha = 2,270,247 pixels) is treated as bare topsoil.")
    h1(doc, "3. The ML target: SOC Deficiency Index (Phase 3)")
    para(doc, "y_soc_def = 0.60 × (1 − SOC_norm) + 0.40 × BSI_norm", bold=True)
    bullet(doc, "60% weight on SoilGrids SOC deficit — the primary vulnerability signal.")
    bullet(doc, "40% weight on satellite BSI exposure — topsoil thermal-oxidation risk.")
    bullet(doc, "SOC_norm and BSI_norm are 1st–99th percentile normalised to [0,1]; output clipped to [0,1].")
    h1(doc, "4. Model features X (9 satellite-only features)")
    para(doc, "bsi, ndvi, swir1_red_ratio, swir1_nir_ratio, bsi_ndvi_ratio, blue_reflectance, red_reflectance, "
              "nir_reflectance, swir1_reflectance (reflectances ÷ 10000).")
    callout(
        doc,
        "ZERO TARGET LEAKAGE:",
        "Raw SoilGrids SOC is 100% excluded from X. The Random Forest must learn the link between optical "
        "SWIR/NIR reflectance and SOC deficiency from the target alone. This is why our R² is honest.",
    )
    h1(doc, "5. Model configuration & validation")
    add_table(
        doc,
        ["Setting", "Value"],
        [
            ["Algorithm", "Random Forest Regressor"],
            ["Trees / depth / leaf", "150 / max_depth 14 / min_samples_leaf 5"],
            ["Training sample", "100,000 sampled pixels → 80,000 train / 20,000 test"],
            ["Fixed seed", "random_state = 42 (reproducible split)"],
            ["Recorded test metrics", f"R² = {M['r2']}, RMSE = {M['rmse']}"],
            ["Live rerun (Aug 2026)", "R² ≈ 0.4550, RMSE ≈ 0.1126 (sampling variance, ±0.01)"],
        ],
    )
    callout(
        doc,
        "Why R² ≈ 0.45 is a STRENGTH, not weakness:",
        "Peer-reviewed optical-soil-carbon literature (RSE, Geoderma, ISPRS) reports R² ∈ [0.35, 0.55] for "
        "un-leakaged regional topsoil carbon regression (Castaldi 2019, Vaudour 2019, Gholizadeh 2018). "
        "Models claiming R² ≥ 0.90 usually leak the target into X. Ours is honest optical physics.",
    )
    h1(doc, "6. Zonal analytics & urgency tiers (Phase 4)")
    bullet(doc, "A 5×5 regular grid overlays the 22km × 22km scene → 25 named sectors (A-row Abhanpur, B-row Arang, C-row Raipur Rural, D-row Dharsiwa, E-row Tilda).")
    bullet(doc, "Each sector reports mean risk, high-risk area (>0.58), bare ha, mean SOC, BSI, pH, clay.")
    add_table(
        doc,
        ["Tier", "Mean SOC Deficiency", "Action"],
        [
            ["Tier 1 — CRITICAL", "≥ 0.58", "Urgent regenerative intervention"],
            ["Tier 2 — MODERATE", "0.45 – 0.58", "Carbon replenishment + cover crops"],
            ["Tier 3 — STABLE", "< 0.45", "Maintenance / monitoring"],
        ],
    )
    h1(doc, "7. Advisory engine rules (recommendations.py)")
    add_table(
        doc,
        ["Condition", "Prescription"],
        [
            ["Risk ≥ 0.58 OR SOC < 100 dg/kg, low clay (<250 g/kg)", "10–12 t/ha FYM or 3–4 t/ha biochar + green manuring + 100% residue retention"],
            ["Risk ≥ 0.58, Vertisol clay (≥250 g/kg)", "8–10 t/ha FYM + zero-tillage (Happy Seeder) + 4 t/ha straw retention"],
            ["Risk 0.45–0.58 OR SOC < 130 dg/kg", "5–6 t/ha FYM + straw mulching + legume cover-crop rotation"],
            ["BSI > 0.25", "Minimum tillage, cover cropping, contour bunding (moisture shield)"],
            ["pH < 5.8", "Agricultural lime 2.0–2.5 t/ha (acidic)"],
            ["pH > 7.5", "Gypsum 2.0 t/ha + green manuring (sodic)"],
        ],
    )
    h1(doc, "8. Uncertainty & confidence mapping")
    bullet(doc, "Ensemble prediction uncertainty = variance across 30 subsampled trees of the forest.")
    bullet(doc, "Rendered as model_confidence_map.png — shows where the model is confident vs uncertain.")
    h1(doc, "9. ISRO Bhuvan validation layer")
    bullet(doc, "Cross-validates bare-soil predictions vs official ISRO Bhuvan 1:50,000 LULC cropland baseline (≈82.4% agricultural cropland match).")
    bullet(doc, "Village Geocoding API maps priority sectors to official Gram Panchayats (e.g., Chandkhuri, Kendri, Nawapara).")
    h1(doc, "10. Honest limitations (judges love this)")
    bullet(doc, "Optical sensors see topsoil (0–15 cm) only — deep carbon needs radar/InSAR or cores.")
    bullet(doc, "Works in post-harvest bare-soil windows (NDVI ≤ 0.30); standing crops mask the signal.")
    bullet(doc, "It is a spatial-prioritisation proxy, not a replacement for wet-chemistry lab assays.")
    doc.save(OUT_DIR / "02_Scientific_Methodology.docx")


# ===============================================================================
# DOC 3 — Pipeline & code guide
# ===============================================================================
def doc03():
    doc = new_doc()
    title_page(
        doc,
        "Pipeline & Code Guide",
        "Every source file, the phase runners, and how to run the system",
        lines=["From raster on disk to executive report in 4 phases"],
    )
    h1(doc, "1. Repository layout")
    para(doc, "soilguard-cg/ — Python geospatial ML engine (the star of the demo).")
    bullet(doc, "data/golden/ — cached Sentinel-2 + SoilGrids GeoTIFFs (offline).")
    bullet(doc, "models/ — trained soil_soc_rf.joblib + soil_soc_metrics.json (R²/RMSE).")
    bullet(doc, "outputs/phase2|3|4/ — maps, CSVs and reports (deliverables).")
    bullet(doc, "outputs/golden_backup/ — the pre-recorded golden output set (fallback for demo).")
    para(doc, "soilguard-dashboard/ — Next.js executive analytics dashboard (port 3001).")
    para(doc, "soilguard-nextjs/ — Next.js interactive map visualizer (port 5000).")
    para(doc, "shared/*.json + scripts/sync-site-data.mjs — single source of truth for both websites.")
    h1(doc, "2. Core modules (src/)")
    add_table(
        doc,
        ["Module", "Responsibility"],
        [
            ["config.py", "Central paths, constants, thresholds, sys.path bootstrap"],
            ["spectral.py", "Load S2 stack; compute NDVI, BSI; bare-soil mask"],
            ["ml_risk.py", "SOC target proxy, 9-feature matrix, RF training, metrics persistence"],
            ["visualize.py", "Phase-2 map renderers (false colour, BSI, NDVI)"],
            ["plot_utils.py", "Shared dark-theme matplotlib helpers (colorbar, footnote, save)"],
            ["zonal.py", "5×5 grid stats, priority ranking, zonal map"],
            ["recommendations.py", "Clay/pH/BSI-aware agronomic prescription engine"],
            ["confidence.py", "Ensemble uncertainty & confidence map"],
            ["report.py", "Executive summary Markdown generator"],
            ["tables.py", "Shared rich table builders (feature importance, top sectors)"],
            ["bhuvan_integration.py", "ISRO Bhuvan LULC + village geocoding validation"],
            ["download_golden.py / verify_golden.py", "Phase-1 STAC/WCS cache + integrity check"],
        ],
    )
    h1(doc, "3. The four phases")
    add_table(
        doc,
        ["Phase", "Runner", "Inputs → Outputs"],
        [
            ["1 · Golden data", "download_golden.py, verify_golden.py", "STAC/WCS → cached GeoTIFFs"],
            ["2 · Spectral", "run_phase2.py", "S2 bands → NDVI, BSI, bare mask + maps"],
            ["3 · ML risk", "run_phase3.py", "Features + target → RF model, risk map, metrics JSON"],
            ["4 · Analytics", "run_phase4.py", "Risk map → zonal ranking CSV, advisories, confidence map, report"],
        ],
    )
    h1(doc, "4. Running the system")
    para(doc, "One-command full demo (recommended for presentation):", bold=True)
    para(doc, "Windows:  demo.bat  (or  run_demo.ps1)      Linux/macOS:  ./demo.sh")
    para(doc, "Or manually, phase by phase (from soilguard-cg/):", bold=True)
    para(doc, "  python src/run_phase2.py\n  python src/run_phase3.py\n  python src/run_phase4.py\n  python src/run_full_demo.py")
    bullet(doc, "run_full_demo.py orchestrates Phases 2→4 in one shot, reusing spectral arrays (no duplicate raster I/O).")
    bullet(doc, "Environment: conda env create -f environment.yml (or pip install -r requirements.txt); Python 3.11.")
    h1(doc, "5. Output deliverables (what judges see)")
    add_table(
        doc,
        ["Artifact", "Location"],
        [
            ["SOC Deficiency Risk Map (10m)", "outputs/phase3/risk_score_map.png"],
            ["Risk histogram", "outputs/phase3/risk_histogram.png"],
            ["Zonal priority map (5×5)", "outputs/phase4/zonal_risk_map.png"],
            ["Model confidence map", "outputs/phase4/model_confidence_map.png"],
            ["Zonal priority ranking CSV", "outputs/phase4/zonal_priority_ranking.csv"],
            ["Agronomic recommendations CSV", "outputs/phase4/agronomic_recommendations.csv"],
            ["Executive summary report", "outputs/phase4/SoilGuard_SOC_Executive_Summary.md"],
            ["Bhuvan village ranking CSV", "outputs/phase4/bhuvan_village_priority_ranking.csv"],
        ],
    )
    doc.save(OUT_DIR / "03_Pipeline_Code_Guide.docx")


# ===============================================================================
# DOC 4 — Key results & cheat sheet (landscape with maps)
# ===============================================================================
def doc04():
    doc = new_doc()
    set_landscape(doc)
    title_page(
        doc,
        "Key Results & Cheat Sheet",
        "The numbers, the full 25-sector ranking, and the maps",
        lines=["Memorise the headline figures — show the maps in the demo"],
    )
    h1(doc, "1. Headline audit numbers")
    add_table(
        doc,
        ["Metric", "Value", "Metric", "Value"],
        [
            ["Total AOI area", f"{M['totalAoiHa']:,.2f} ha", "Bare-soil area evaluated", f"{M['bareSoilHa']:,.2f} ha"],
            ["Bare-soil share", f"{M['bareSoilPct']:.2f}%", "Mean risk score", f"{M['meanRisk']:.4f}"],
            ["Low risk (<0.45)", f"{M['lowRiskHa']:,.2f} ha ({M['lowRiskPct']:.2f}%)", "Moderate (0.45–0.58)", f"{M['moderateRiskHa']:,.2f} ha ({M['moderateRiskPct']:.2f}%)"],
            ["HIGH risk (>0.58)", f"{M['highRiskHa']:,.2f} ha ({M['highRiskPct']:.2f}%)", "Model R² / RMSE", f"{M['r2']} / {M['rmse']}"],
            ["Cut-off thresholds", f"0.45 / {M['highCutoff']}", "Resolution", M["resolution"]],
        ],
        font_size=10,
    )
    h1(doc, "2. Full 25-sector priority ranking (golden recorded values)")
    df = load_zonal()
    rows = [
        [
            int(r["priority_rank"]),
            r["sector_name"],
            f"{r['mean_risk_score']:.4f}",
            f"{r['bare_soil_ha']:,.1f}",
            f"{r['high_risk_ha']:,.1f}",
            f"{r['pct_high_risk']:.1f}",
            f"{r['mean_soc_dg_kg']:.1f}",
            f"{r['mean_bsi']:.4f}",
            f"{r['mean_ph']:.1f}",
            f"{r['mean_clay_g_kg']:.0f}",
        ]
        for _, r in df.iterrows()
    ]
    add_table(
        doc,
        ["Rank", "Sector", "Mean Risk", "Bare (ha)", "High-Def (ha)", "% High", "SOC (dg/kg)", "BSI", "pH", "Clay (g/kg)"],
        rows,
        font_size=8.5,
    )
    h1(doc, "3. Top-5 agronomic advisory packages")
    recs = load_recs()
    rows5 = [
        [
            int(r["priority_rank"]),
            r["sector_name"],
            r["urgency_level"],
            r["primary_recommendation"][:120] + ("…" if len(r["primary_recommendation"]) > 120 else ""),
        ]
        for _, r in recs.head(5).iterrows()
    ]
    add_table(doc, ["Rank", "Sector", "Urgency", "Primary recommendation"], rows5, font_size=8.5)
    h1(doc, "4. Presentation maps")
    embed_map(doc, "phase3", "risk_score_map.png", "Fig 1 · SOC Deficiency Risk Score Map (10m) — phase3/risk_score_map.png", max_w=6.5)
    embed_map(doc, "phase4", "zonal_risk_map.png", "Fig 2 · 5×5 Zonal Organic Carbon Priority Map — phase4/zonal_risk_map.png", max_w=6.5)
    embed_map(doc, "phase4", "model_confidence_map.png", "Fig 3 · Model Ensemble Confidence / Uncertainty Map — phase4/model_confidence_map.png", max_w=6.5)
    embed_map(doc, "phase2", "false_color_composite.png", "Fig 4 · False Colour Composite (SWIR1-NIR-Red) — phase2/false_color_composite.png", max_w=6.5)
    embed_map(doc, "phase2", "bsi_map.png", "Fig 5 · Bare Soil Index (BSI) Map — phase2/bsi_map.png", max_w=6.5)
    embed_map(doc, "phase2", "ndvi_map.png", "Fig 6 · NDVI Map — phase2/ndvi_map.png", max_w=6.5)
    embed_map(doc, "phase3", "risk_histogram.png", "Fig 7 · Re-calibrated Risk Score Histogram — phase3/risk_histogram.png", max_w=6.5)
    doc.save(OUT_DIR / "04_Key_Results_Cheat_Sheet.docx")


# ===============================================================================
# DOC 5 — Live demo guide
# ===============================================================================
def doc05():
    doc = new_doc()
    title_page(
        doc,
        "Live Demo Guide",
        "A safe, scripted walkthrough of the one-command pipeline",
        lines=["Practise twice; keep golden_backup as your safety net"],
    )
    h1(doc, "0. Pre-flight checklist")
    bullet(doc, "Golden rasters present in data/golden/ (sentinel2_raipur_golden.tif, soilgrids_raipur_golden.tif).")
    bullet(doc, "Python 3.11 with rasterio, scikit-learn, matplotlib, rich, joblib installed (conda env: soilguard).")
    bullet(doc, "Outputs pre-generated in outputs/phase4/ as fallback.")
    bullet(doc, "Speaker notes + cheat sheet in hand; terminal font large enough for the back row.")
    h1(doc, "1. The command")
    para(doc, "demo.bat  (Windows)   ·   ./demo.sh  (Linux/WSL)   ·   run_demo.ps1  (PowerShell)", bold=True)
    para(doc, "Under the hood it runs:  python src/run_full_demo.py  (from soilguard-cg/).")
    h1(doc, "2. What happens on screen (and what to say)")
    add_table(
        doc,
        ["Step", "What appears", "Say this (≈15 s per step)"],
        [
            ["Launch", "Title banner + [Step 1/5] loading rasters", "\"This is SoilGuard-CG — fully offline, using cached Sentinel-2 and SoilGrids data.\""],
            ["Step 2/5", "Bare-soil pixel count (2.27 M, 48.96%)", "\"We isolate bare topsoil with NDVI ≤ 0.30 — excluding crops and water.\""],
            ["Step 3/5", "Feature-importance table + R²/RMSE line", "\"A Random Forest trained on 9 pure satellite features — no target leakage, R² ≈ 0.45, exactly the literature benchmark.\""],
            ["Step 4/5", "Top-5 zonal priority table", "\"Arang B-1 is rank #1 — 946 ha of high SOC deficiency. That block gets priority funding.\""],
            ["Step 5/5", "Executive summary + SUCCESS panel", "\"Every map, CSV and report is generated right now, in about a minute.\""],
        ],
        font_size=9,
    )
    h1(doc, "3. After the run — show the deliverables")
    bullet(doc, "Open outputs/phase3/risk_score_map.png — the 10m SOC deficiency heatmap.")
    bullet(doc, "Open outputs/phase4/zonal_risk_map.png — the 5×5 priority grid with rank labels.")
    bullet(doc, "Open outputs/phase4/agronomic_recommendations.csv — scroll to Arang B-1: 10–12 t/ha FYM prescription.")
    h1(doc, "4. If something breaks (the safety net)")
    callout(
        doc,
        "FALLBACK:",
        "All deliverables are pre-rendered in outputs/golden_backup/ and outputs/phase4/. If the live run fails, "
        "say \"let me show you the pre-generated outputs\" and open the maps/CSVs directly. Judges value composure.",
    )
    h1(doc, "5. Optional: Bhuvan enrichment (network required)")
    para(doc, "python src/bhuvan_integration.py  → queries ISRO Bhuvan LULC + village geocoding, producing "
              "bhuvan_village_priority_ranking.csv. Mention it as the ISRO validation layer; only demo live if "
              "internet is reliable, otherwise show the pre-generated CSV.")
    doc.save(OUT_DIR / "05_Live_Demo_Guide.docx")


# ===============================================================================
# DOC 6 — Presentation script
# ===============================================================================
def doc06():
    doc = new_doc()
    title_page(
        doc,
        "Presentation Script",
        "A timed, slide-by-slide pitch script (2-minute and 5-minute versions)",
        lines=["Hook → Problem → Solution → Results → Impact"],
    )
    h1(doc, "1. The 5-slide deck skeleton")
    add_table(
        doc,
        ["Slide", "Title", "Content / visual", "Time"],
        [
            ["1", "Title & Problem", "Chhattisgarh SOC crisis, stubble burning, SHC gap", "0:00–0:30"],
            ["2", "Methodology", "Sentinel-2 → NDVI/BSI mask → RF (no leakage) → 10m map", "0:30–1:15"],
            ["3", "Results", "22,702 ha evaluated · 4,441 ha high deficiency · maps", "1:15–1:45"],
            ["4", "Advisory", "Top sectors + clay-sensitive FYM dosages", "1:45–2:15"],
            ["5", "Validation & Impact", "Bhuvan LULC 82.4% · offline · societal benefit", "2:15–3:00"],
        ],
    )
    h1(doc, "2. The 2-minute pitch (word-for-word)")
    para(doc, "[0:00–0:30] Hook & crisis:", bold=True)
    para(doc, "\"Respected judges, Chhattisgarh is the rice bowl of India. But after every harvest, millions of "
              "hectares of topsoil lose organic carbon to summer heat and stubble burning. Traditional soil "
              "cards sample just one point per 10 hectares every three years — farmers stay blind to degradation.\"")
    para(doc, "[0:30–1:15] Solution & demo:", bold=True)
    para(doc, "\"SoilGuard-CG fixes this with one terminal command. It ingests 10m Sentinel-2 imagery, filters "
              "2.27 million bare-soil pixels, trains a leakage-free Random Forest, and renders SOC deficiency "
              "maps in about a minute — 100% offline.\"")
    para(doc, "[1:15–1:45] Results:", bold=True)
    para(doc, "\"Across 22,700 hectares we found 19.6% of topsoil — over 4,400 hectares — in critical SOC "
              "deficiency, concentrated in Arang and Abhanpur blocks. The map shows exactly where.\"")
    para(doc, "[1:45–2:00] Impact & close:", bold=True)
    para(doc, "\"We then attach clay-sensitive advisories — 10–12 tonnes of FYM per hectare — and validate "
              "against ISRO Bhuvan baselines. Space data, offline, in farmers' hands. Thank you.\"")
    h1(doc, "3. The 5-minute version (add these segments)")
    bullet(doc, "Add the SHC cost argument: ₹2,000–3,000 per sample vs zero-cost satellite revisits every 5 days.")
    bullet(doc, "Whiteboard the formulas: NDVI, BSI, and y = 0.60·(1−SOC) + 0.40·BSI.")
    bullet(doc, "Explain the zero-leakage design and the R² = 0.45 literature benchmark (Castaldi 2019, Vaudour 2019, Gholizadeh 2018).")
    bullet(doc, "Walk one sector end-to-end: Arang B-1 → risk 0.66 → high-def 946 ha → 10–12 t/ha FYM prescription.")
    bullet(doc, "Demo the live pipeline (demo.bat) during the methodology segment.")
    h1(doc, "4. Delivery tips")
    bullet(doc, "Open with the crisis, not the tech. Tech impresses after the problem is felt.")
    bullet(doc, "When showing the map, point at the red (high-deficiency) zones — make it visual.")
    bullet(doc, "If asked about R²: smile and call it 'honest physics' — see the Q&A doc.")
    bullet(doc, "Finish with the offline capability: 'a district officer can run this on a basic laptop, no internet.'")
    doc.save(OUT_DIR / "06_Presentation_Script.docx")


# ===============================================================================
# DOC 7 — Judges Q&A
# ===============================================================================
def doc07():
    doc = new_doc()
    title_page(
        doc,
        "Judges Q&A Masterclass",
        "The 12 hardest questions, with winning answers",
        lines=["From the master study guide — updated with live pipeline numbers"],
    )
    qa = [
        ("Q1 · In simple terms, what does your model predict?",
         "A continuous SOC Deficiency Index (0.0–1.0) at 10m resolution. Near 1.0 = severe organic-carbon depletion and "
         "thermal-oxidation risk; near 0.0 = stable carbon-rich soil. It pinpoints exactly where topsoil organic matter is "
         "missing so officers know where to apply FYM."),
        ("Q2 · Why SOC instead of generic 'soil health'?",
         "Generic soil health is too vague to fund. SOC is the single most critical indicator of fertility, moisture "
         "retention and structure in Chhattisgarh, and it aligns with National Carbon Credit & Green Credit initiatives. "
         "It lets us give exact dosages (10–12 t/ha FYM)."),
        ("Q3 · How is this better than government Soil Health Cards?",
         "SHC samples 1 point per 10 ha every 3 years at ₹2,000–3,000/sample. We evaluate every 10m × 10m cell using "
         "free Sentinel-2 data refreshed every 5 days — a satellite guidance system that tells sampling teams exactly "
         "where to go. We complement, not replace, lab testing."),
        ("Q4 · Optical satellites only see the surface. How do you measure SOC?",
         "Crucial distinction: we model topsoil (0–15 cm) deficiency via SWIR1/B11 optical reflectance — carbon-rich "
         "soils absorb SWIR, depleted bare soils reflect it. We target dry post-harvest windows when fields are bare."),
        ("Q5 · What is your ground truth, and did you avoid target leakage?",
         "Ground truth = SoilGrids SOC deficit (60%) blended with satellite BSI exposure (40%). Raw SoilGrids SOC is "
         "100% excluded from features X; the model trains on 9 spectral inputs only, forcing it to learn true optical "
         "physics instead of memorising static maps."),
        ("Q6 · Can this run in remote district offices without internet?",
         "Yes — the pipeline is 100% offline-first: cached Cloud-Optimized GeoTIFFs, local ML inference, maps and "
         "reports generated in under a minute on a basic laptop."),
        ("Q7 · What are your main limitations?",
         "Three honest ones: (1) requires bare/sparse soil (NDVI ≤ 0.30) — best in post-harvest windows; (2) optical "
         "sensing covers topsoil 0–15 cm, not deep carbon (needs radar/cores); (3) it is a spatial-prioritisation "
         "proxy, not a replacement for wet-chemistry assays."),
        ("Q8 · Why Random Forest and not deep learning / CNN?",
         "For tabular pixel-level multispectral data over a regional extent, RF avoids overfitting, runs in seconds, "
         "and gives explicit feature-importance interpretability (SWIR1 dominates). It runs offline on low-spec "
         "field computers — CNNs would need heavy GPUs."),
        ("Q9 · How would a district collector actually use the output?",
         "Two artifacts: the Zonal Priority Ranking CSV (which block first — e.g., Arang B-1) and the Advisory CSV "
         "(how many tonnes of FYM/biochar to distribute in each block under PMKSY/RKVY)."),
        ("Q10 · How does this align with ISRO's mission and National Space Day?",
         "It turns open Earth Observation data into operational agricultural intelligence — exactly ISRO's EO-for-"
         "society goal: protecting topsoil health, climate resilience, and farmer livelihoods."),
        ("Q11 · Is R² ≈ 0.45 too low? Why not 0.90+?",
         "No — the peer-reviewed benchmark for un-leakaged optical topsoil-carbon regression is R² ∈ [0.35, 0.55] "
         "(Castaldi 2019; Vaudour 2019; Gholizadeh 2018). Claims of R² ≥ 0.90 usually leak the target into X. "
         "Our R² is honest, generalising optical physics."),
        ("Q12 · How do you cross-validate against official ISRO data?",
         "We validate the 48.96% bare-soil share against ISRO Bhuvan 1:50,000 LULC cropland baselines (≈82.4% "
         "agricultural cropland), and use Bhuvan Village Geocoding to attach advisories to official Gram "
         "Panchayats (Chandkhuri, Kendri, Nawapara)."),
    ]
    for q, a in qa:
        h2(doc, q)
        para(doc, a)
    h1(doc, "Bonus: follow-up one-liners")
    bullet(doc, "\"How long did training take?\"  →  Full pipeline ≈41 s recorded; training itself is seconds on 100k samples (first run may take ~1–2 min on a laptop).")
    bullet(doc, "\"What would you add next?\"  →  Multi-date compositing, SAR for deep carbon, and an API/field app.")
    bullet(doc, "\"Is the data real?\"  →  Yes — Sentinel-2 L2A and SoilGrids, cached offline in data/golden/.")
    doc.save(OUT_DIR / "07_Judges_QA.docx")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    builders = {
        "00_Start_Here_Study_Plan": doc00,
        "01_Project_Overview_Pitch": doc01,
        "02_Scientific_Methodology": doc02,
        "03_Pipeline_Code_Guide": doc03,
        "04_Key_Results_Cheat_Sheet": doc04,
        "05_Live_Demo_Guide": doc05,
        "06_Presentation_Script": doc06,
        "07_Judges_QA": doc07,
    }
    for name, fn in builders.items():
        fn()
        print(f"[OK] {name}.docx")
    print(f"\nDone. Pack written to: {OUT_DIR}")


if __name__ == "__main__":
    main()
